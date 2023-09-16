from docx import Document 
from docx.shared import Inches
import os
from pathlib import Path

r,c=9,5
document = Document()
tables = document.add_table(rows=r, cols=c)
def pypart(r,c):
    #x=os.listdir()
    try:
        result = [os.path.join(dp, f) for dp, dn, filenames in os.walk("C:\\Users\\RAJU\\Downloads\\i words") for f in filenames if os.path.splitext(f)[1] == '.png']
        a=0
        for i in range(0, r):
    
            for j in range(0, c):
                print(" i  " , i , " j " , j)
                #tables = document.add_table(rows=3, cols=3)
                p = tables.rows[i].cells[j].add_paragraph()
                r = p.add_run()
                if len(result[a]) > 2:
                    r.add_picture(result[a],width=Inches(1.4), height=Inches(1.1))
                #document.save('addImage.docx')
                a=a+1	
    except:
        print("error occured in except")
    finally:
        document.save('addImage.docx')

# Driver Code
#n = 4
pypart(r,c)
#document.save('addImage.docx')
    
    
